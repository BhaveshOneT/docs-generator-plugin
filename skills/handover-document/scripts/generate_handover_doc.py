#!/usr/bin/env python3
"""
Generate a branded One Thousand handover document (DOCX) from structured JSON content.

Usage:
    python generate_handover_doc.py \
        --content /tmp/handover_content.json \
        --logo-dir /path/to/logos/ \
        --output /path/to/output/handover.docx \
        [--arch-diagram /tmp/arch_diagram.png]
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
except ImportError:
    Image = None

# Brand colors
SHARP_GREEN = RGBColor(0x19, 0xA9, 0x60)
DARK_TEXT = RGBColor(0x2F, 0x2F, 0x2F)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "19A960"
CODE_BG = "F5F5F5"
CODE_BORDER = "DDDDDD"

# Font names
HEADING_FONT = "Amsi Pro Narw Black"
BODY_FONT = "Akkurat LL"
CODE_FONT = "Courier New"

# Fallback fonts
HEADING_FONT_FALLBACK = "Arial Black"
BODY_FONT_FALLBACK = "Calibri"


def set_font(run, font_name, size, color, bold=False, italic=False):
    """Set font properties on a run with fallback."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    # Set East Asian font fallback
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)


def add_cover_page(doc, content, logo_dir):
    """Create the green branded cover page."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    # Green background using a full-page rectangle
    # We'll use a colored table cell to simulate the background
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)

    # Set cell background to green
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_HEADER_BG}" w:val="clear"/>')
    tcPr.append(shading)

    # Set cell dimensions to fill page
    cell.width = Cm(21.0)
    tr = table.rows[0]._element
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="16840" w:hRule="exact"/>')
    trPr.append(trHeight)

    # Remove table borders
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Remove default cell margins/padding
    cell_margins = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'<w:top w:w="0" w:type="dxa"/>'
        f'<w:left w:w="0" w:type="dxa"/>'
        f'<w:bottom w:w="0" w:type="dxa"/>'
        f'<w:right w:w="0" w:type="dxa"/>'
        f'</w:tblCellMar>'
    )
    tblPr.append(cell_margins)

    # Add content to cell
    # Logo
    logo_path = os.path.join(logo_dir, "onethousand-icon-limeonblack-rounded.png")
    if os.path.exists(logo_path):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(80)
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(1.2))

    # "handover" text
    language = content.get("language", "en")
    display_word = "übergabe" if language == "de" else "handover"

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(display_word)
    set_font(run, HEADING_FONT, 48, WHITE, bold=True)

    # Project name
    project = content.get("project", {})
    project_name = project.get("name", "Project")
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(project_name)
    set_font(run, HEADING_FONT, 24, WHITE)

    # Client attribution
    client_name = project.get("client", "")
    if client_name:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(15)
        p.paragraph_format.space_after = Pt(0)
        attribution = f"{client_name} × One Thousand"
        run = p.add_run(attribution)
        set_font(run, HEADING_FONT, 18, WHITE)

    # Handover date
    handover_date = project.get("handover_date", "")
    if handover_date:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(30)
        run = p.add_run(handover_date)
        set_font(run, BODY_FONT, 14, WHITE)

    # Contacts
    contacts = content.get("contacts", {})
    ot_team = contacts.get("ot_team", [])
    client_team = contacts.get("client_team", [])

    if ot_team or client_team:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(40)

        # Participants label
        label_text = "Beteiligte:" if language == "de" else "Participants:"
        run = p.add_run(label_text)
        set_font(run, BODY_FONT, 12, WHITE, bold=True)
        run.underline = True

        # Client team
        if client_team:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(0)
            names = sorted([c.get("name", "") for c in client_team])
            for i, name in enumerate(names):
                if i > 0:
                    run = p.add_run("\n")
                    set_font(run, BODY_FONT, 11, WHITE)
                run = p.add_run(name)
                set_font(run, BODY_FONT, 11, WHITE)

        # Spacer
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)

        # OT team
        if ot_team:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            names = sorted([c.get("name", "") for c in ot_team])
            for i, name in enumerate(names):
                if i > 0:
                    run = p.add_run("\n")
                    set_font(run, BODY_FONT, 11, WHITE)
                run = p.add_run(name)
                set_font(run, BODY_FONT, 11, WHITE)


def add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run._element.append(parse_xml(f'<w:br {nsdecls("w")} w:type="page"/>'))


def add_toc(doc, language="en"):
    """Add table of contents."""
    # Reset margins for content pages
    section = doc.add_section()
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    toc_title = "Inhaltsverzeichnis" if language == "de" else "Table of Contents"
    p = doc.add_paragraph()
    run = p.add_run(toc_title)
    set_font(run, HEADING_FONT, 28, SHARP_GREEN)
    p.paragraph_format.space_after = Pt(20)

    # TOC field code
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fld_char_begin)

    run = p.add_run()
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run._element.append(instr)

    run = p.add_run()
    fld_char_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run._element.append(fld_char_separate)

    run = p.add_run("[Update table of contents to populate]")
    set_font(run, BODY_FONT, 11, DARK_TEXT)

    run = p.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._element.append(fld_char_end)


def add_heading1(doc, text):
    """Add a green H1 heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, HEADING_FONT, 16, SHARP_GREEN, bold=True)

    # Set as Heading 1 for TOC
    p.style = doc.styles['Heading 1']
    # Override the style font
    run.font.name = HEADING_FONT
    run.font.size = Pt(16)
    run.font.color.rgb = SHARP_GREEN
    run.font.bold = True

    return p


def add_heading2(doc, text):
    """Add a green H2 heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, HEADING_FONT, 14, SHARP_GREEN, bold=True)

    p.style = doc.styles['Heading 2']
    run.font.name = HEADING_FONT
    run.font.size = Pt(14)
    run.font.color.rgb = SHARP_GREEN
    run.font.bold = True

    return p


def add_heading3(doc, text):
    """Add a green H3 heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, HEADING_FONT, 12, SHARP_GREEN, bold=True)

    p.style = doc.styles['Heading 3']
    run.font.name = HEADING_FONT
    run.font.size = Pt(12)
    run.font.color.rgb = SHARP_GREEN
    run.font.bold = True

    return p


def add_body_text(doc, text):
    """Add body text with inline formatting support."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _add_formatted_runs(p, text)
    return p


def _add_formatted_runs(paragraph, text):
    """Parse markdown inline formatting and add runs."""
    # Process bold and italic
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            set_font(run, BODY_FONT, 11, DARK_TEXT, bold=True, italic=True)
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_font(run, BODY_FONT, 11, DARK_TEXT, bold=True)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            set_font(run, BODY_FONT, 11, DARK_TEXT, italic=True)
        elif part:
            run = paragraph.add_run(part)
            set_font(run, BODY_FONT, 11, DARK_TEXT)


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)

    # Clear default run and add formatted
    for run in p.runs:
        run.text = ""
    _add_formatted_runs(p, text)
    return p


def add_numbered_item(doc, text, level=0):
    """Add a numbered list item."""
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)

    for run in p.runs:
        run.text = ""
    _add_formatted_runs(p, text)
    return p


def add_code_block(doc, code_text):
    """Add a code block with grey background."""
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)

        # Grey background shading
        pPr = p._element.get_or_add_pPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>')
        pPr.append(shading)

        run = p.add_run(line)
        set_font(run, CODE_FONT, 10, DARK_TEXT)

    return p


def add_markdown_table(doc, header_row, data_rows):
    """Add a formatted table with green headers."""
    num_cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(header_row):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header.strip().replace('**', ''))
        set_font(run, BODY_FONT, 10, WHITE, bold=True)

        # Green background
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_HEADER_BG}" w:val="clear"/>')
        tcPr.append(shading)

    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.cell(row_idx + 1, col_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                _add_formatted_runs(p, cell_text.strip())

    return table


def parse_markdown_content(doc, markdown_text, arch_diagram_path=None):
    """Parse markdown content and add it to the document."""
    lines = markdown_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code block handling
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table handling
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            # Check if separator row
            if all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            # Check if next line is still table
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
                i += 1
                continue
            else:
                # End of table, render it
                if len(table_rows) > 1:
                    add_markdown_table(doc, table_rows[0], table_rows[1:])
                table_rows = []
                i += 1
                continue

        # Flush any pending table
        if table_rows and not line.strip().startswith('|'):
            if len(table_rows) > 1:
                add_markdown_table(doc, table_rows[0], table_rows[1:])
            table_rows = []

        # Headings
        if line.startswith('## '):
            add_heading2(doc, line[3:].strip())
            i += 1
            continue
        elif line.startswith('### '):
            add_heading3(doc, line[4:].strip())
            i += 1
            continue

        # Image embedding (for architecture diagram)
        if line.strip().startswith('!['):
            match = re.match(r'!\[.*?\]\((.*?)\)', line.strip())
            if match:
                img_path = match.group(1)
                if img_path.startswith('data:image'):
                    # Base64 image - skip for now, handled separately
                    pass
                elif os.path.exists(img_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(6))
            elif arch_diagram_path and os.path.exists(arch_diagram_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(arch_diagram_path, width=Inches(6))
            i += 1
            continue

        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            level = (len(line) - len(line.lstrip())) // 2
            add_bullet(doc, text, level)
            i += 1
            continue

        # Numbered lists
        match = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if match:
            level = len(match.group(1)) // 2
            text = match.group(3)
            add_numbered_item(doc, text, level)
            i += 1
            continue

        # Empty line
        if line.strip() == '':
            i += 1
            continue

        # Regular paragraph
        add_body_text(doc, line)
        i += 1

    # Flush remaining table
    if table_rows and len(table_rows) > 1:
        add_markdown_table(doc, table_rows[0], table_rows[1:])


def add_page_numbers(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Page number field
        run = p.add_run()
        fld_char = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fld_char)

        run = p.add_run()
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run._element.append(instr)

        run = p.add_run()
        fld_char = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run._element.append(fld_char)

        run = p.add_run("1")
        set_font(run, BODY_FONT, 10, DARK_TEXT)

        run = p.add_run()
        fld_char = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._element.append(fld_char)


def generate_handover_doc(content_path, logo_dir, output_path, arch_diagram_path=None):
    """Main generation function."""
    # Load content
    with open(content_path, 'r', encoding='utf-8') as f:
        content = json.load(f)

    language = content.get("language", "en")

    # Create document
    doc = Document()

    # Cover page
    add_cover_page(doc, content, logo_dir)

    # Page break after cover
    add_page_break(doc)

    # Table of contents
    add_toc(doc, language)

    # Page break after TOC
    add_page_break(doc)

    # Content sections
    sections = content.get("sections", [])
    for idx, section in enumerate(sections):
        section_id = section.get("id", "")
        title = section.get("title", f"Section {idx + 1}")
        section_content = section.get("content", "")

        # Add H1 heading
        add_heading1(doc, title)

        # Parse and add markdown content
        # Pass arch diagram path for the architecture section
        diagram_path = arch_diagram_path if section_id == "architecture" else None
        parse_markdown_content(doc, section_content, diagram_path)

        # Page break between major sections (except last)
        if idx < len(sections) - 1:
            add_page_break(doc)

    # Page numbers
    add_page_numbers(doc)

    # Save
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    print(f"Handover document generated: {output_path}")
    print(f"Sections: {len(sections)}")
    print(f"Language: {language}")

    # Validate
    try:
        validation_doc = Document(output_path)
        print(f"Validation: {len(validation_doc.paragraphs)} paragraphs, "
              f"{len(validation_doc.tables)} tables")
    except Exception as e:
        print(f"WARNING: Validation failed: {e}")


def main():
    parser = argparse.ArgumentParser(description="Generate One Thousand branded handover document")
    parser.add_argument('--content', required=True, help='Path to content JSON file')
    parser.add_argument('--logo-dir', required=True, help='Path to logos directory')
    parser.add_argument('--output', required=True, help='Output DOCX file path')
    parser.add_argument('--arch-diagram', default=None, help='Path to architecture diagram PNG')

    args = parser.parse_args()

    if not os.path.exists(args.content):
        print(f"ERROR: Content file not found: {args.content}")
        sys.exit(1)

    if args.arch_diagram and not os.path.exists(args.arch_diagram):
        print(f"WARNING: Architecture diagram not found: {args.arch_diagram}")
        args.arch_diagram = None

    generate_handover_doc(args.content, args.logo_dir, args.output, args.arch_diagram)


if __name__ == '__main__':
    main()
